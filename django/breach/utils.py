import requests
from .models import *
from docx import Document
import tempfile
from io import BytesIO
import os
import uuid
from docx2pdf import convert
from io import BytesIO
import pythoncom
from .serializers import *
import datetime
def fetch_and_store_breach_data(email, user):
    url = f"https://api.xposedornot.com/v1/breach-analytics?email={email}"
    response = requests.get(url)
    if response.status_code != 200:
        return {"error": "API call failed"}

    data = response.json()
    breach_metrics = data.get('BreachMetrics') or {}
    risk_list = breach_metrics.get('risk') or [{}]
    risk_info = risk_list[0] if risk_list else {}
    pastes_summary = data.get('PastesSummary') or {}
    breaches_summary = data.get('BreachesSummary') or {}
    exposed_breaches = data.get('ExposedBreaches') or {}

    # Create or update EmailBreachRecord
    record, created = EmailBreachRecord.objects.get_or_create(email=email, defaults={"user": user})
    
    # If the record already exists, update its fields and delete related data
    if not created:
        record.user = user
        record.breach_risk_label = risk_info.get('risk_label')
        record.breach_risk_score = risk_info.get('risk_score')
        record.paste_count = pastes_summary.get('cnt')
        record.last_paste_timestamp = pastes_summary.get('tmpstmp')
        record.save()

        # Clear old related data
        record.breach_summaries.all().delete()
        record.exposed_breaches.all().delete()
        record.industries.all().delete()
        record.password_strengths.all().delete()
        record.yearly_stats.all().delete()
        for cat in record.exposed_data.all():
            cat.items.all().delete()
        record.exposed_data.all().delete()
    else:
        record.breach_risk_label = risk_info.get('risk_label')
        record.breach_risk_score = risk_info.get('risk_score')
        record.paste_count = pastes_summary.get('cnt')
        record.last_paste_timestamp = pastes_summary.get('tmpstmp')
        record.save()

    # Breach Summary
    BreachSummary.objects.create(
        email_record=record,
        site=breaches_summary.get('site')
    )

    # Exposed Breaches
    for breach in exposed_breaches.get('breaches_details', []):
        try:
            ExposedBreach.objects.create(
                email_record=record,
                breach=breach.get('breach'),
                details=breach.get('details'),
                domain=breach.get('domain'),
                industry=breach.get('industry'),
                logo=breach.get('logo'),
                password_risk=breach.get('password_risk'),
                references=breach.get('references'),
                searchable=breach.get('searchable'),
                verified=breach.get('verified'),
                xposed_data=breach.get('xposed_data'),
                xposed_date=breach.get('xposed_date'),
                xposed_records=breach.get('xposed_records')
            )
        except Exception as e:
            print(f"Error creating ExposedBreach: {e}")

    # Industry
    for industry_list in breach_metrics.get('industry', []):
        for industry_entry in industry_list:
            try:
                industry_code = industry_entry[0] if len(industry_entry) > 0 else None
                count = industry_entry[1] if len(industry_entry) > 1 else None
                BreachMetricIndustry.objects.create(
                    email_record=record,
                    industry_code=industry_code,
                    count=count
                )
            except Exception as e:
                print(f"Error creating BreachMetricIndustry: {e}")

    # Password Strength
    for ps in breach_metrics.get('passwords_strength', []):
        try:
            PasswordStrength.objects.create(
                email_record=record,
                easy_to_crack=ps.get('EasyToCrack'),
                plain_text=ps.get('PlainText'),
                strong_hash=ps.get('StrongHash'),
                unknown=ps.get('Unknown')
            )
        except Exception as e:
            print(f"Error creating PasswordStrength: {e}")

    # Year-wise stats
    for year_stat in breach_metrics.get('yearwise_details', []):
        for key, count in year_stat.items():
            try:
                year = int(key.lstrip("y")) if key else None
                YearlyBreachStats.objects.create(
                    email_record=record,
                    year=year,
                    count=count
                )
            except Exception as e:
                print(f"Error creating YearlyBreachStats: {e}")

    # Exposed Data Types
    for data_entry in breach_metrics.get('xposed_data', []):
        for level2 in data_entry.get('children', []):
            try:
                cat = ExposedDataCategory.objects.create(
                    email_record=record,
                    level2_name=level2.get('name')
                )
                for item in level2.get('children', []):
                    try:
                        ExposedDataItem.objects.create(
                            category=cat,
                            group=item.get('group'),
                            name=item.get('name'),
                            value=item.get('value')
                        )
                    except Exception as e:
                        print(f"Error creating ExposedDataItem: {e}")
            except Exception as e:
                print(f"Error creating ExposedDataCategory: {e}")

    return {"success": True, "email": record.email}


# breach/utils.py



def create_breach_report_pdf(user, email_record):
    x=datetime.datetime.now()
    exposed_data=[]
    try:
        for i in email_record.exposed_data.all():
            exposed_data.append(i['level2_name'])
    except:
        exposed_data.append('NO leaks found')
    # Initialize COM for this thread (Windows only)
    pythoncom.CoInitialize()

    try:
        # Create temp folder
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, f"{uuid.uuid4()}.docx")
            pdf_path = docx_path.replace(".docx", ".pdf")

            # 1. Create .docx file
            doc = Document()
            doc.add_heading('Dark Web Monitoring Report', 0)

            doc.add_paragraph(f'Application User: {user.username}')
            doc.add_paragraph(f'Report Generated: {x.year} {x.strftime("%B")} {x.day}')
            doc.add_paragraph(f'Email : {user.email}')
            doc.add_paragraph(f'Risk Score: {email_record.breach_risk_score} ({email_record.breach_risk_label})')
            doc.add_paragraph(f'Paste Count: {email_record.paste_count}')
            doc.add_paragraph(f'Last Paste Timestamp: {email_record.last_paste_timestamp}')
            
           

            doc.add_heading('Exposed Data', level=1)
            for i in email_record.exposed_data.all():
                doc.add_paragraph(f'    : {i.level2_name}')
    

            doc.add_heading('Exposed Breaches', level=1)
            try:
                for exposed_breach in email_record.exposed_breaches.all():
                    doc.add_paragraph(f'Breach: {exposed_breach.breach}')
                    doc.add_paragraph(f'Domain: {exposed_breach.domain}')
                    doc.add_paragraph(f'Industry: {exposed_breach.industry}')
                    doc.add_paragraph(f'details:{exposed_breach.details}')
                
            except:
                doc.add_paragraph(f'*No breach found')


            doc.add_heading('Security Recommendations', level=0)
            for r in [
                "Change compromised passwords immediately.",
                "Enable two-factor authentication.",
                "Use a password manager.",
                "Monitor financial accounts for unusual activity.",
                "Be alert to phishing attempts."
            ]:
                doc.add_paragraph(f'• {r}')

            doc.save(docx_path)

            # 2. Convert to PDF
            convert(docx_path, pdf_path)

            # 3. Read PDF to memory
            with open(pdf_path, "rb") as f:
                pdf_buffer = BytesIO(f.read())

        return pdf_buffer
    finally:
        pythoncom.CoUninitialize()  # Clean up COM